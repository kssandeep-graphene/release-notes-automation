import streamlit as st
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import re
from typing import List, Dict, Tuple
import json
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException

class LinkExtractor:
    def _get_selenium_driver(self):
        """Initializes a headless Chrome WebDriver."""
        try:
            options = webdriver.ChromeOptions()
            options.add_argument('--headless')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
            
            # Suppress console logs from webdriver_manager
            os.environ['WDM_LOG_LEVEL'] = '0'
            
            service = ChromeService(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            return driver
        except Exception as e:
            st.error(f"❌ Failed to initialize Selenium WebDriver: {e}")
            st.warning("Please ensure Google Chrome is installed on your system.")
            return None

    def extract_links_from_url(self, url: str) -> Dict:
        """Return a dict with 'links', 'toc_tree', 'page_title'. Uses JS to crawl shadow DOM for anchors and TOC hierarchy."""
        st.info(f"🚀 Processing URL with Selenium: {url}")
        driver = self._get_selenium_driver()
        
        if not driver:
            return {'links': [], 'toc_tree': [], 'page_title': f"Error: WebDriver initialization failed"}
        
        try:
            driver.get(url)
            
            # Wait for the page title to be meaningful, indicating the page has likely loaded.
            wait = WebDriverWait(driver, 25) # 25-second timeout
            wait.until(lambda d: d.title and "Help And Training Community" not in d.title and "Login" not in d.title)
            
            page_title = driver.title
            st.success(f"📄 Page loaded successfully. Title: {page_title}")
            
            #####################################################################
            # NEW APPROACH: use JavaScript to walk the DOM (including shadow DOM) 
            # and collect all anchors, because Salesforce renders the article body
            # inside Lightning-web-components shadow roots that are invisible to
            # BeautifulSoup / driver.page_source.
            #####################################################################

            JS_CRAWL_ANCHORS = """
            const anchors = [];
            function walk(node){
                if(!node) return;
                try{
                    if(node.nodeType===1){
                        if(node.tagName==='A' && node.href){
                            anchors.push({href: node.href, text: node.innerText.trim()});
                        }
                        // regular children
                        for(const child of node.children){ walk(child); }
                        // shadow DOM
                        if(node.shadowRoot){
                            for(const srChild of node.shadowRoot.children){ walk(srChild); }
                        }
                    }
                }catch(e){ /* ignore cross-origin errors */ }
            }
            walk(document.body);
            return anchors;
            """

            raw_links = driver.execute_script(JS_CRAWL_ANCHORS)
            st.info(f"🔎 JS crawler returned {len(raw_links)} anchors total")

            links = []
            for item in raw_links:
                href = item.get('href', '')
                text = item.get('text', '').strip()
                if not text or not href:
                    continue

                # Only keep release-notes/articleView targets or descriptive text
                if 'release-notes' not in href:
                    continue

                # Filter obvious navigation words
                if any(word in text.lower() for word in ['home','login','support','contact','privacy','terms','footer','navigation','refresh','print']):
                    continue

                links.append({
                    'text': text,
                    'url': href,
                    'original_href': href
                })

            st.success(f"✅ After filtering: {len(links)} release-note anchors")
            
            # Show the first few links for debugging
            for i, link in enumerate(links[:3]):
                st.info(f"Link {i+1}: '{link['text'][:60]}...' -> {link['url'][:60]}...")
            
            # Remove duplicates
            seen_urls = set()
            unique_links = []
            for link in links:
                if link['url'] not in seen_urls:
                    seen_urls.add(link['url'])
                    unique_links.append(link)

            # ------------------------------------------------------------------
            # ALSO build hierarchical tree from the Table of Contents sidebar
            # ------------------------------------------------------------------
            JS_TOC_TREE = """
            function serialize(node){
               const item = {text: (node.innerText||'').trim(), url: null, children: []};
               const a = node.querySelector(':scope > a[href]');
               if(a){ item.text = a.innerText.trim(); item.url = a.href; }
               const childrenLis = node.querySelectorAll(':scope > ul > li');
               childrenLis.forEach(li => item.children.push(serialize(li)));
               return item;
            }
            // Locate the TOC container – Salesforce uses 'table-of-content' class
            let toc = document.querySelector('[class*="table-of-content"] ul');
            if(!toc){ return []; }
            const out = [];
            toc.querySelectorAll(':scope > li').forEach(li => out.push(serialize(li)));
            return out;
            """

            toc_tree = driver.execute_script(JS_TOC_TREE)
            st.info(f"📑 TOC tree captured with {len(toc_tree)} top-level nodes")

            return {
                'links': unique_links,
                'toc_tree': toc_tree,
                'page_title': page_title
            }

        except TimeoutException:
            st.error(f"❌ Timed out waiting for page to load: {url}")
            st.info("The page may be too slow or protected. Could not extract content.")
            return {'links': [], 'toc_tree': [], 'page_title': f"Error: Timeout for {url}"}
        except Exception as e:
            st.error(f"❌ An error occurred during extraction: {str(e)}")
            import traceback
            st.error(f"Full traceback: {traceback.format_exc()}")
            return {'links': [], 'toc_tree': [], 'page_title': f"Error: {url}"}
        finally:
            if driver:
                driver.quit()

class DocxManager:
    def __init__(self):
        self.document = None
    
    def create_document(self, title: str, extracted_data: list) -> io.BytesIO:
        """Create a new DOCX document with extracted links"""
        try:
            # Create a new document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, level=0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add creation date
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            date_run.italic = True
            
            # Add a line break
            doc.add_paragraph()
            
            # Add summary
            summary_heading = doc.add_heading('Summary', level=1)
            total_links = sum(len(data['links']) for data in extracted_data)
            total_sources = len(extracted_data)
            
            summary_para = doc.add_paragraph()
            summary_para.add_run("Total Sources Processed: ").bold = True
            summary_para.add_run(f"{total_sources}\n")
            summary_para.add_run(f"Total Links Extracted: ").bold = True
            summary_para.add_run(f"{total_links}")
            
            doc.add_page_break()
            
            # Add content for each URL
            for i, url_data in enumerate(extracted_data, 1):
                source_url = url_data['source_url']
                page_title = url_data['page_title']
                links = url_data.get('links', [])
                toc_tree = url_data.get('toc_tree', [])
                
                # Add header for this source
                source_heading = doc.add_heading(f"{i}. {page_title}", level=1)
                
                # Add source URL
                source_para = doc.add_paragraph()
                source_para.add_run("Source: ").bold = True
                self._add_hyperlink(source_para, source_url, source_url)
                
                # Add link count
                count_para = doc.add_paragraph()
                count_para.add_run("Links Found: ").bold = True
                count_para.add_run(str(len(links)))
                
                doc.add_paragraph()  # Add space
                
                # Write hierarchical TOC if available
                if toc_tree:
                    def write_node(node, level=0):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.25 * level)
                        if node.get('url'):
                            self._add_hyperlink(p, node['url'], node['text'])
                        else:
                            p.add_run(node['text']).bold = True
                        for child in node.get('children', []):
                            write_node(child, level + 1)
                    for top in toc_tree:
                        write_node(top, 0)
                elif links:
                    # Fallback flat list
                    for j, link in enumerate(links, 1):
                        link_para = doc.add_paragraph()
                        link_para.add_run(f"{j}. ").bold = True
                        self._add_hyperlink(link_para, link['url'], link['text'])
                        doc.add_paragraph()
                
                # Add separator between sources (except for the last one)
                if i < len(extracted_data):
                    separator_para = doc.add_paragraph("=" * 80)
                    separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
            
            # Save to BytesIO buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as error:
            st.error(f"An error occurred creating document: {error}")
            return None
    
    def _add_hyperlink(self, paragraph, url, text):
        """Add a working hyperlink using the most reliable method"""
        try:
            # Validate URL
            if not url or not url.startswith(('http://', 'https://')):
                return paragraph.add_run(text)
            
            # Use the built-in hyperlink functionality with error handling
            from docx.oxml.shared import qn, oxml_parser
            from docx.oxml.ns import nsdecls, qn as qname
            
            # Get the document's part
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Create hyperlink using raw XML - this is the most reliable method
            hyperlink_xml = f'''<w:hyperlink {nsdecls('w', 'r')} r:id="{r_id}"><w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r></w:hyperlink>'''
            
            hyperlink_element = oxml_parser.parse_xml(hyperlink_xml)
            paragraph._element.append(hyperlink_element)
            
            return hyperlink_element
            
        except Exception as e:
            # Simpler fallback that at least shows the URL
            try:
                # Try to create a basic functional hyperlink
                run = paragraph.add_run()
                
                # Create hyperlink XML element directly
                from docx.oxml.shared import qn
                part = paragraph.part
                r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                
                # Create hyperlink element
                hyperlink = paragraph._element.makeelement(qn('w:hyperlink'))
                hyperlink.set(qn('r:id'), r_id)
                
                # Create run element
                new_run = paragraph._element.makeelement(qn('w:r'))
                
                # Add text
                text_elem = paragraph._element.makeelement(qn('w:t'))
                text_elem.text = text
                new_run.append(text_elem)
                
                # Add run properties for styling
                rPr = paragraph._element.makeelement(qn('w:rPr'))
                color = paragraph._element.makeelement(qn('w:color'))
                color.set(qn('w:val'), '0563C1')
                underline = paragraph._element.makeelement(qn('w:u'))
                underline.set(qn('w:val'), 'single')
                rPr.append(color)
                rPr.append(underline)
                new_run.insert(0, rPr)
                
                hyperlink.append(new_run)
                paragraph._element.append(hyperlink)
                
                return hyperlink
                
            except:
                # Final fallback: Show both text and URL so user can copy the URL
                run = paragraph.add_run(text)
                run.font.color.rgb = RGBColor(5, 99, 193)  # Standard hyperlink blue
                run.underline = True
                
                # Add line break and the actual URL
                paragraph.add_run("\n    ")
                url_run = paragraph.add_run(url)
                url_run.font.color.rgb = RGBColor(5, 99, 193)
                url_run.underline = True
                url_run.font.size = url_run.font.size * 0.9 if url_run.font.size else None
                
                return run

def main():
    st.set_page_config(
        page_title="Link Extractor to DOCX",
        page_icon="🔗",
        layout="wide"
    )
    
    st.title("🔗 Link Extractor to DOCX")
    st.markdown("Extract links from multiple URLs and organize them in a Word document")
    
    # Initialize session state with proper defaults
    if 'extracted_data' not in st.session_state:
        st.session_state.extracted_data = []
    if 'processing' not in st.session_state:
        st.session_state.processing = False
    
    # Sidebar for instructions
    with st.sidebar:
        st.header("📋 Instructions")
        st.markdown("""
        1. **Add URLs**: Enter URLs (one per line) in the text area
        2. **Extract Links**: Click the 'Extract Links' button
        3. **Review**: Check the extracted links in the preview
        4. **Export**: Click 'Download DOCX' to get the Word document
        
        **Note**: The DOCX file can be opened with Microsoft Word or Google Docs.
        """)
        
        st.header("✨ Features")
        st.markdown("""
        - **Smart Extraction**: Handles JavaScript-heavy sites like Salesforce
        - **Structured Output**: Organized by source with headers
        - **DOCX Format**: Compatible with Word and Google Docs
        - **Clickable Links**: URLs are styled and ready to copy
        - **No Setup Required**: No API credentials needed
        """)
        
        st.header("🎯 Perfect For")
        st.markdown("""
        - Release notes documentation
        - API documentation links
        - Changelog compilations
        - Reference link collections
        """)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📝 Input URLs")
        
        # URL input with a unique key to avoid session state issues
        url_input = st.text_area(
            "Enter URLs (one per line):",
            height=200,
            placeholder="https://example.com/release-notes\nhttps://docs.example.com/api\nhttps://changelog.example.com",
            key="url_input_area"
        )
        
        # Extract button
        if st.button("🔍 Extract Links", disabled=st.session_state.processing, key="extract_button"):
            if url_input.strip():
                urls = [url.strip() for url in url_input.split('\n') if url.strip()]
                
                if urls:
                    st.session_state.processing = True
                    st.session_state.extracted_data = []
                    
                    extractor = LinkExtractor()
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for i, url in enumerate(urls):
                        status_text.text(f"Processing: {url}")
                        
                        extracted_data = extractor.extract_links_from_url(url)
                        
                        st.session_state.extracted_data.append({
                            'source_url': url,
                            'page_title': extracted_data['page_title'],
                            'links': extracted_data['links'],
                            'toc_tree': extracted_data['toc_tree']
                        })
                        
                        progress_bar.progress((i + 1) / len(urls))
                    
                    status_text.text("✅ Extraction complete!")
                    st.session_state.processing = False
                    st.rerun()
                else:
                    st.warning("Please enter at least one valid URL")
            else:
                st.warning("Please enter URLs to extract links from")
    
    with col2:
        st.header("📊 Extracted Links Preview")
        
        if st.session_state.extracted_data:
            # Summary
            total_links = sum(len(data['links']) for data in st.session_state.extracted_data)
            st.metric("Total Links Found", total_links)
            st.metric("URLs Processed", len(st.session_state.extracted_data))
            
            # Preview of extracted data
            with st.expander("🔍 Preview Extracted Links", expanded=True):
                for data in st.session_state.extracted_data:
                    st.subheader(f"📄 {data['page_title']}")
                    st.caption(f"Source: {data['source_url']}")
                    
                    if data['links']:
                        for i, link in enumerate(data['links'][:5], 1):  # Show first 5 links
                            st.write(f"{i}. [{link['text']}]({link['url']})")
                        
                        if len(data['links']) > 5:
                            st.caption(f"... and {len(data['links']) - 5} more links")
                    else:
                        st.warning("No links found on this page")
                    
                    st.divider()
            
            # DOCX export section
            st.header("📄 Export to DOCX")
            
            doc_title = st.text_input(
                "Document Title:",
                value="Extracted Links Report",
                key="doc_title_input"
            )
            
            if st.button("📄 Create & Download DOCX", key="download_button"):
                with st.spinner("Creating DOCX document..."):
                    try:
                        docx_manager = DocxManager()
                        docx_buffer = docx_manager.create_document(doc_title, st.session_state.extracted_data)
                        
                        if docx_buffer:
                            # Generate filename with timestamp
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"{doc_title.replace(' ', '_')}_{timestamp}.docx"
                            
                            st.success("✅ DOCX document created successfully!")
                            
                            # Download button with explicit MIME type
                            st.download_button(
                                label="⬇️ Download DOCX File",
                                data=docx_buffer.getvalue(),
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_docx_button"
                            )
                            
                            st.info("💡 **Tips**:")
                            st.info("- Right-click links in the document to copy URLs")
                            st.info("- Open with Microsoft Word or upload to Google Docs")
                            st.info("- All URLs are formatted as blue, underlined text")
                            st.balloons()
                        else:
                            st.error("Failed to create DOCX document")
                    except Exception as e:
                        st.error(f"Error creating document: {str(e)}")
        else:
            st.info("👆 Enter URLs and click 'Extract Links' to see preview")

if __name__ == "__main__":
    main() 